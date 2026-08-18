
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── CÀI ĐẶT TRANG ───────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.0)
section.top_margin  = section.bottom_margin = Inches(1.0)

# ─── FONT MẶC ĐỊNH ───────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc, text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = color
    p.runs[0].font.name = 'Times New Roman'
    return p

def add_para(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11.5)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x4C, 0x97)
    return p

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def make_table_header(table, headers, fill='1F497D'):
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        shade_cell(hdr_cells[i], fill)

def add_table_row(table, values, shade=None):
    row = table.add_row()
    for i, v in enumerate(values):
        row.cells[i].text = str(v)
        for para in row.cells[i].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
        if shade:
            shade_cell(row.cells[i], shade)
    return row

# ══════════════════════════════════════════════════════════════
# TRANG BÌA
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TRƯỜNG ĐẠI HỌC TÔN ĐỨC THẮNG')
run.bold = True; run.font.size = Pt(14); run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
run.bold = True; run.font.size = Pt(13); run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BÁO CÁO TIỂU LUẬN')
run.bold = True; run.font.size = Pt(20); run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LẬP TRÌNH CƠ SỞ DỮ LIỆU')
run.bold = True; run.font.size = Pt(15); run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GIẢI THÍCH MÃ NGUỒN')
run.bold = True; run.font.size = Pt(26); run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HỆ THỐNG QUẢN LÝ CỬA HÀNG PHÂN PHỐI MÁY TÍNH')
run.bold = True; run.font.size = Pt(16); run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

info = [
    ('Sinh viên thực hiện', 'Nguyễn Yến Nhi'),
    ('MSSV', 'DH23IM01'),
    ('Giảng viên hướng dẫn', 'ThS. Phạm Chí Công'),
    ('Năm học', '2025 – 2026'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}: ')
    r1.bold = True; r1.font.size = Pt(13); r1.font.name = 'Times New Roman'
    r2 = p.add_run(value)
    r2.font.size = Pt(13); r2.font.name = 'Times New Roman'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# MỤC LỤC (viết tay)
# ══════════════════════════════════════════════════════════════
add_heading(doc, 'MỤC LỤC', 1)
toc_items = [
    ('I.',   'KIẾN TRÚC N-LAYER — TỔNG QUAN VÀ THỨ TỰ THIẾT KẾ'),
    ('II.',  'LỚP DTO — Data Transfer Objects'),
    ('III.', 'LỚP DAL — Data Access Layer'),
    ('IV.',  'LỚP BUS — Business Logic Layer'),
    ('V.',   'LỚP GUI — Windows Forms (Ứng dụng nhân viên)'),
    ('VI.',  'LỚP WEB — ASP.NET MVC (Website khách hàng)'),
    ('VII.', 'SO SÁNH GUI VÀ WEB — Các điều kiện bổ sung so với BUS'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{num}  ')
    r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
    r2 = p.add_run(title)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHẦN I — KIẾN TRÚC N-LAYER
# ══════════════════════════════════════════════════════════════
add_heading(doc, 'PHẦN I — KIẾN TRÚC N-LAYER: TỔNG QUAN VÀ THỨ TỰ THIẾT KẾ', 1)

add_para(doc,
    'Hệ thống được xây dựng theo mô hình N-Layer (đa tầng) — một kiến trúc '
    'phần mềm chia ứng dụng thành các tầng (layer) độc lập, mỗi tầng chỉ được '
    'phép giao tiếp với tầng ngay bên dưới nó. Mục tiêu chính là tách biệt '
    'trách nhiệm, giảm sự phụ thuộc chéo, dễ bảo trì và mở rộng.')
doc.add_paragraph()

add_heading(doc, '1.1 Sơ đồ kiến trúc', 2)
add_code(doc, 'GUI_WinForm  ──┐')
add_code(doc, '               ├──▶  BUS  ──▶  DAL  ──▶  DTO')
add_code(doc, 'Web_MVC      ──┘')
doc.add_paragraph()

add_para(doc,
    'Luồng phụ thuộc chỉ đi một chiều từ giao diện → nghiệp vụ → truy xuất '
    'dữ liệu → đối tượng truyền dữ liệu. Tuyệt đối không được phép '
    'gọi DAL trực tiếp từ GUI hoặc Web.', bold=False)
doc.add_paragraph()

add_heading(doc, '1.2 Ý nghĩa từng tầng', 2)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Tầng (Layer)', 'Tên Project', 'Vai Trò'])
rows_data = [
    ('DTO', 'DTO_HTQLCuaHangLaptop', 'Định nghĩa cấu trúc dữ liệu — chỉ chứa thuộc tính (properties). Không có logic.'),
    ('DAL', 'DAL_HTQLCuaHangLaptop', 'Truy xuất CSDL SQL Server — thực hiện CRUD thuần túy bằng ADO.NET. Không validate, không tính toán.'),
    ('BUS', 'BUS_HTQLCuaHangLaptop', 'Xử lý nghiệp vụ — kiểm tra dữ liệu đầu vào, áp dụng quy tắc kinh doanh, điều phối DAL.'),
    ('GUI', 'GUI_HTQLCuaHangLaptop', 'Giao diện Windows Forms dành cho nhân viên — hiển thị dữ liệu, thu thập input, phân quyền.'),
    ('Web', 'Website (ASP.NET MVC)', 'Giao diện website dành cho khách hàng — đặt hàng, xem sản phẩm, quản lý tài khoản.'),
]
shades = ['DEEAF1', 'FFFFFF', 'DEEAF1', 'FFFFFF', 'DEEAF1']
for (t, p_name, r), shade in zip(rows_data, shades):
    add_table_row(tbl, [t, p_name, r], shade=shade)
doc.add_paragraph()

add_heading(doc, '1.3 Thứ tự thiết kế và lập trình khuyến nghị', 2)
steps = [
    ('Bước 1 — CSDL SQL Server',
     'Tạo toàn bộ 22 bảng với ràng buộc khoá, index, CHECK constraint. '
     'File QuanLyCuaHangLaptop.sql đã có đầy đủ.'),
    ('Bước 2 — DTO',
     'Mỗi bảng → 1 class DTO. Viết trước để tất cả các tầng sau đều dùng được. '
     'Không có dependency nào — class thuần chứa property.'),
    ('Bước 3 — DBConnect + DAL',
     'Viết class DBConnect kết nối CSDL, sau đó từng class DAL_X cho mỗi bảng. '
     'Bắt đầu từ bảng nền tảng: VaiTro, TaiKhoanNV, NhanVien, HangSanXuat...'),
    ('Bước 4 — BUS',
     'Viết BUS từ module đơn giản đến phức tạp: BUS_SanPham → BUS_KhachHang '
     '→ BUS_TaiKhoan → BUS_KhuyenMai → BUS_DonHang → BUS_HauMai → BUS_BaoCao.'),
    ('Bước 5 — GUI WinForm',
     'Bắt đầu từ FormDangNhap, FormMain (phân quyền), sau đó từng form nghiệp vụ.'),
    ('Bước 6 — Web MVC',
     'HomeController → SanPhamController → TaiKhoanController → GioHangController → ...'),
    ('Bước 7 — Kiểm thử tích hợp',
     'Test từng luồng: đăng nhập → tạo đơn → khuyến mãi → bảo hành → báo cáo.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    r1 = p.add_run(f'✅ {title}: ')
    r1.bold = True; r1.font.size = Pt(11.5); r1.font.name = 'Times New Roman'
    r2 = p.add_run(desc)
    r2.font.size = Pt(11.5); r2.font.name = 'Times New Roman'
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHẦN II — LỚP DTO
# ══════════════════════════════════════════════════════════════
add_heading(doc, 'PHẦN II — LỚP DTO (Data Transfer Objects)', 1)

add_para(doc,
    'DTO (Data Transfer Object) là các class đơn giản chỉ chứa thuộc tính '
    '(properties), không có phương thức nghiệp vụ, không import namespace DAL '
    'hay BUS. Mỗi DTO tương ứng 1:1 với một bảng trong CSDL, đóng vai trò '
    '"hộp chứa dữ liệu" để truyền thông tin giữa các tầng.')
doc.add_paragraph()

add_heading(doc, '2.1 Quy tắc thiết kế DTO', 2)
rules = [
    'Tên class: tiền tố DTO_ + tên bảng (ví dụ: DTO_SanPham, DTO_KhachHang).',
    'Chỉ khai báo public properties — không có constructor đặc biệt, không có logic.',
    'Dùng kiểu dữ liệu nullable (DateTime?, decimal?) cho cột có thể NULL trong CSDL.',
    'Dùng decimal cho tiền tệ — tránh float/double gây sai số tài chính.',
    'Cột bool IsDeleted phục vụ xóa mềm (soft delete) — không xóa vật lý bản ghi master.',
    'Audit columns: NgayTao, NgayCapNhat, NguoiTao, NguoiCapNhat được thêm vào các bảng nghiệp vụ.',
]
for r in rules:
    add_bullet(doc, r)
doc.add_paragraph()

add_heading(doc, '2.2 Danh sách 23 DTO và ý nghĩa', 2)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Tên DTO', 'Bảng CSDL', 'Ý nghĩa / Thuộc tính chính'])
dto_data = [
    ('DTO_VaiTro', 'VaiTro', 'Lưu mã và tên vai trò (VT001-VT005). Không có IsDeleted — phân quyền cố định.'),
    ('DTO_TaiKhoanNV', 'TaiKhoanNV', 'Tài khoản đăng nhập nhân viên. Có MaVaiTro (FK), MatKhau (lưu dạng SHA-256 hash), TrangThai (Hoạt Động / Khóa).'),
    ('DTO_TaiKhoanKH', 'TaiKhoanKH', 'Tài khoản đăng nhập khách hàng website. Tương tự TaiKhoanNV nhưng có MaKH thay MaNV.'),
    ('DTO_LichSuDangNhap', 'LichSuDangNhap', 'Ghi lại mỗi lần đăng nhập NV. Có MaTK, DiaChiIP, TrangThai (Thành Công / Thất Bại), NgayDangNhap.'),
    ('DTO_NhanVien', 'NhanVien', 'Hồ sơ nhân viên: MaNV, TenNV, CCCD, SĐT, NgaySinh, GioiTinh, DiaChi, ChucVu, IsDeleted.'),
    ('DTO_KhachHang', 'KhachHang', 'Hồ sơ khách hàng chung. Có LoaiKH ("Lẻ" / "Sỉ"), Email, SDT, DiaChi, IsDeleted.'),
    ('DTO_KhachHangLe', 'KhachHangLe', 'Mở rộng KhachHang cho KH lẻ. Có LaHSSV (bool) xác định học sinh/sinh viên để áp dụng KM Back To School.'),
    ('DTO_KhachHangSi', 'KhachHangSi', 'Mở rộng KhachHang cho KH sỉ (doanh nghiệp). Có TenCongTy, MaSoThue.'),
    ('DTO_HangSanXuat', 'HangSanXuat', 'Hãng sản xuất laptop/phụ kiện: MaHang, TenHang, QuocGia, IsDeleted.'),
    ('DTO_LoaiSanPham', 'LoaiSanPham', 'Loại sản phẩm: MaLoaiSP, MaHang (FK), TenLoai, DanhMuc ("Laptop"/"Chuột"/"Bàn Phím"), ThoiGianBaoHanh (tháng), GiaBanGoc, IsDeleted.'),
    ('DTO_CauHinh', 'CauHinh', 'Cấu hình kỹ thuật của LoaiSanPham: CPU, RAM, ổ cứng, màn hình (TenThuocTinh, GiaTri).'),
    ('DTO_SanPham', 'SanPham', 'Sản phẩm vật lý: MaSerialSP (PK, unique), MaPhieuNhap (FK), MaLoaiSP (FK), NgayNhap, NgaySX, TrangThai ("Trong Kho"/"Đã Bán"/"Bảo Hành"/"Lỗi"/"Đổi Trả"), IsDeleted.'),
    ('DTO_NhaCungCap', 'NhaCungCap', 'Nhà cung cấp: MaNCC, TenNCC, DiaChi, Email, SDT, IsDeleted.'),
    ('DTO_PhieuNhap', 'PhieuNhap', 'Phiếu nhập hàng: MaPhieuNhap, MaNV, MaNCC, NgayNhap, TongTien, TrangThai ("Chờ Xác Nhận"/"Đã Nhập"/"Huỷ").'),
    ('DTO_ChiTietPhieuNhap', 'ChiTietPhieuNhap', 'Chi tiết phiếu nhập: MaPhieuNhap, MaLoaiSP, SoLuong, DonGiaNhap.'),
    ('DTO_KhuyenMai', 'KhuyenMai', 'Chương trình KM: MaKM, TenKM, DoiTuong, NgayBatDau, NgayKetThuc, DieuKien, SLToiThieu, MucGiamSP (%), MucGiamDH (%).'),
    ('DTO_HopDong', 'HopDong', 'Hợp đồng với KH sỉ: MaHD, MaKH, MaNV, NgayKy, NgayHetHan, GiaTriHopDong, TrangThai ("Hiệu Lực"/"Hết Hạn"/"Huỷ").'),
    ('DTO_DonHang', 'DonHang', 'Đơn hàng: MaDH, MaNV, MaKH, MaKM?, MaHD?, TongTien, TienSauGiam, PhuongThucThanhToan, TrangThai ("Chờ Xử Lý"/"Đang Giao"/"Hoàn Thành"/"Huỷ"), NgayDat.'),
    ('DTO_ChiTietDonHang', 'ChiTietDonHang', 'Chi tiết đơn hàng: MaDH, MaSerialSP (UNIQUE — mỗi serial chỉ bán 1 lần), GiaBan, PhanTramGiam.'),
    ('DTO_PhieuBaoHanh', 'PhieuBaoHanh', 'Phiếu bảo hành: MaPBH, MaDH, MaKH, MaSerialSP, LoaiBH ("Cửa Hàng"/"Hãng"), NgayBatDau, NgayKetThuc, TrangThai, KetQua.'),
    ('DTO_PhieuDoiTra', 'PhieuDoiTra', 'Phiếu đổi trả: MaPhieuDT, MaDH, MaKH, MaSerialSP (UNIQUE), LyDo, LoaiXuLy ("Đổi Máy"/"Hoàn Tiền"/"Từ Chối"), TrangThai.'),
    ('DTO_DonKhieuNai', 'DonKhieuNai', 'Đơn khiếu nại: MaDonKN, MaDH (phải "Hoàn Thành"), MaKH, NoiDung, TrangThai ("Đang Xử Lý"/"Đã Giải Quyết"/"Từ Chối").'),
]
for i, (dto, bang, y_nghia) in enumerate(dto_data):
    shade = 'DEEAF1' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [dto, bang, y_nghia], shade=shade)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHẦN III — LỚP DAL
# ══════════════════════════════════════════════════════════════
add_heading(doc, 'PHẦN III — LỚP DAL (Data Access Layer)', 1)

add_para(doc,
    'DAL là tầng truy xuất dữ liệu — lớp duy nhất trong hệ thống được phép '
    'viết câu SQL và kết nối SQL Server. Mỗi class DAL kế thừa từ DBConnect '
    'và thực hiện CRUD thuần cho một bảng/nhóm bảng liên quan.')
doc.add_paragraph()

add_heading(doc, '3.1 Class DBConnect — Nền tảng kết nối', 2)
add_para(doc,
    'DBConnect là class cơ sở mà mọi DAL kế thừa. Nó đọc connection string '
    '"QuanLyCuaHangLaptop" từ App.config (WinForm) / appsettings.json (Web), '
    'tạo đối tượng SqlConnection dùng chung, và cung cấp hai phương thức '
    'tiện ích OpenConnection() và CloseConnection().')
add_code(doc, 'protected SqlConnection _conn;')
add_code(doc, 'public DBConnect() {')
add_code(doc, '    string connStr = ConfigurationManager')
add_code(doc, '        .ConnectionStrings["QuanLyCuaHangLaptop"].ConnectionString;')
add_code(doc, '    _conn = new SqlConnection(connStr);')
add_code(doc, '}')
doc.add_paragraph()

add_heading(doc, '3.2 Quy tắc lập trình DAL', 2)
dal_rules = [
    'Luôn gọi _conn.Open() trước khi thực thi câu lệnh SQL, gọi CloseConnection() trong khối finally để đảm bảo kết nối luôn được đóng dù có lỗi.',
    'Dùng SqlParameter để truyền tham số — tuyệt đối không ghép chuỗi SQL trực tiếp (SQL Injection Prevention).',
    'Giá trị trả về: DataTable cho danh sách, DTO_X cho đơn lẻ, bool cho CUD (Create/Update/Delete).',
    'Không chứa logic nghiệp vụ (validate, tính toán) — chỉ CRUD thuần.',
    'Không import namespace BUS hay GUI.',
]
for r in dal_rules:
    add_bullet(doc, r)
doc.add_paragraph()

add_heading(doc, '3.3 Danh sách 24 class DAL và chức năng', 2)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Tên Class DAL', 'Bảng Phụ Trách', 'Các Phương Thức Chính'])
dal_data = [
    ('DBConnect', '(Base class)', 'OpenConnection(), CloseConnection(). Đọc connection string từ config.'),
    ('DAL_VaiTro', 'VaiTro', 'DSTatCaVaiTro(), DSTheoMaVaiTro(maVT), CapNhatVaiTro(dto).'),
    ('DAL_TaiKhoanNV', 'TaiKhoanNV', 'DSTatCaTaiKhoanNV(), DSTheoMaTK(maTK), DSTheoTenDangNhap(ten), ThemTaiKhoanNV(dto), CapNhatTaiKhoanNV(dto), CapNhatTrangThai(maTK, trangThai), LayMaTKNVMoiNhat(), LayDanhSachNVChuaCoTaiKhoan().'),
    ('DAL_TaiKhoanKH', 'TaiKhoanKH', 'DSTatCaTaiKhoanKH(), DSTheoMaTK(maTK), DSTheoTenDangNhap(ten), DSTheoMaKH(maKH), ThemTaiKhoanKH(dto), CapNhatTaiKhoanKH(dto), CapNhatTrangThai(maTK, tt), LayMaTKKHMoiNhat().'),
    ('DAL_LichSuDangNhap', 'LichSuDangNhap', 'DSTatCaLichSuDangNhap(), DSTheoMaTK(maTK), ThemLichSuDangNhap(dto).'),
    ('DAL_NhanVien', 'NhanVien', 'DSTatCaNhanVien(), DSTheoMaNV(maNV), ThemNhanVien(dto), CapNhatNhanVien(dto), XoaMemNhanVien(maNV).'),
    ('DAL_KhachHang', 'KhachHang', 'DSTheoMaKH(maKH), DSTheoEmail(email), ThemKhachHang(dto), CapNhatKhachHang(dto), XoaMemKhachHang(maKH), LayMaKHLeMoiNhat(), LayMaKHSiMoiNhat().'),
    ('DAL_KhachHangLe', 'KhachHangLe', 'DSTheoMaKHLe(maKH), ThemKhachHangLe(dto), CapNhatKhachHangLe(dto).'),
    ('DAL_KhachHangSi', 'KhachHangSi', 'DSTheoMaKHSi(maKH), ThemKhachHangSi(dto), CapNhatKhachHangSi(dto).'),
    ('DAL_HangSanXuat', 'HangSanXuat', 'DSTatCaHSX(), DSTheoMaHSX(maHang), ThemHSX(dto), UpdateHangSanXuat(dto), XoaMemHSX(maHang).'),
    ('DAL_LoaiSanPham', 'LoaiSanPham', 'DSLoaiSP(), DSLoaiSPTheoHang(maHang), DSLoaiSPTheoDanhMuc(danhMuc), TimLoaiSP(maLoaiSP), ThemLoaiSP(dto), CapNhatLoaiSP(dto), XoaMemLoaiSP(maLoaiSP).'),
    ('DAL_CauHinh', 'CauHinh', 'DSTatCaCauHinh(), DSCauHinhTheoLoaiSP(maLoaiSP), DSTheoMaCauHinh(maCH), ThemCauHinh(dto), CapNhatCauHinh(dto), XoaCauHinh(maCH).'),
    ('DAL_SanPham', 'SanPham', 'DSTatCaSanPham(), DSTheoTrangThai(tt), DSTheoLoaiSP(maLoaiSP), DSTheoMaSerialSP(serial), DSTheoPhieuNhap(maPN), ThemSanPham(dto), CapNhatSanPham(dto), CapNhatTrangThai(serial, tt), XoaMemSanPham(serial).'),
    ('DAL_NhaCungCap', 'NhaCungCap', 'DSTatCaNCC(), DSTheoMaNCC(maNCC), ThemNCC(dto), CapNhatNCC(dto), XoaMemNCC(maNCC).'),
    ('DAL_PhieuNhap', 'PhieuNhap', 'DSTatCaPhieuNhap(), DSTheoMaPN(maPN), ThemPhieuNhap(dto), CapNhatPhieuNhap(dto), CapNhatTrangThai(maPN, tt).'),
    ('DAL_ChiTietPhieuNhap', 'ChiTietPhieuNhap', 'DSTheoPhieuNhap(maPN), ThemChiTietPhieuNhap(dto), XoaChiTietTheoPhieuNhap(maPN).'),
    ('DAL_KhuyenMai', 'KhuyenMai', 'DSTatCaKhuyenMai(), DSTrongThoiGianHieuLuc(ngay), DSTheoMaKM(maKM), ThemKhuyenMai(dto), CapNhatKhuyenMai(dto), XoaKhuyenMai(maKM).'),
    ('DAL_HopDong', 'HopDong', 'DSTatCaHopDong(), DSTheoMaHD(maHD), DSTheoKhachHang(maKH), DSHieuLuc(), ThemHopDong(dto), CapNhatHopDong(dto), CapNhatTrangThai(maHD, tt).'),
    ('DAL_DonHang', 'DonHang', 'DSTatCaDonHang(), DSTheoMaDH(maDH), DSTheoKhachHang(maKH), DSTheoTrangThai(tt), CapNhatTrangThai(maDH, tt), LayMaDHMoiNhat().'),
    ('DAL_ChiTietDonHang', 'ChiTietDonHang', 'DSTheoDonHang(maDH), DSTheoMaSerialSP(serial), DSTatCaChiTietDonHang(), DSChiTietCoThongTinSanPham(maDH).'),
    ('DAL_PhieuBaoHanh', 'PhieuBaoHanh', 'DSTatCaPhieuBaoHanh(), DSTheoKhachHang(maKH), DSTheoMaSerial(serial), DSTheoMaPhieuBaoHanh(maPBH), ThemPhieuBaoHanh(dto), CapNhatTrangThai(maPBH, tt), CapNhatKetQua(maPBH, ketQua), CapNhatLyDoLoi(maPBH, lyDo).'),
    ('DAL_PhieuDoiTra', 'PhieuDoiTra', 'DSTatCaPhieuDoiTra(), DSTheoKhachHang(maKH), DSTheoMaSerial(serial), DSTheoMaPhieuDoiTra(maPDT), ThemPhieuDoiTra(dto), CapNhatTrangThai(maPDT, tt), CapNhatLoaiXuLy(maPDT, loaiXL).'),
    ('DAL_DonKhieuNai', 'DonKhieuNai', 'DSTatCaDonKhieuNai(), DSTheoMaDonKhieuNai(maDKN), DSTheoMaSerial(serial), ThemDonKhieuNai(dto), CapNhatTrangThai(maDKN, tt).'),
]
for i, (name, bang, methods) in enumerate(dal_data):
    shade = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [name, bang, methods], shade=shade)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHẦN IV — LỚP BUS
# ══════════════════════════════════════════════════════════════
add_heading(doc, 'PHẦN IV — LỚP BUS (Business Logic Layer)', 1)

add_para(doc,
    'BUS là tầng nghiệp vụ — trái tim của hệ thống. BUS nhận yêu cầu từ GUI/Web, '
    'kiểm tra tính hợp lệ của dữ liệu đầu vào, áp dụng quy tắc kinh doanh, '
    'sau đó gọi DAL để thực thi. BUS không có câu SQL, không import SqlClient.')
doc.add_paragraph()

# ── BUS_SanPham ──────────────────────────────────
add_heading(doc, '4.1 BUS_SanPham — Quản lý danh mục sản phẩm', 2)
add_para(doc, 'Phụ trách 4 thực thể: HangSanXuat, LoaiSanPham, CauHinh, SanPham (serial vật lý).')

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Nhóm', 'Phương Thức', 'Kiểm Tra / Điều Kiện'])
bus_sp_data = [
    ('HangSanXuat', 'LayDanhSachHSX()', 'Không có điều kiện đặc biệt — lấy tất cả HSX chưa xóa mềm.'),
    ('HangSanXuat', 'ThemHSX(hsx)', 'TenHang không rỗng, ≤100 ký tự. Không trùng tên với hãng đang hoạt động.'),
    ('HangSanXuat', 'CapNhatHSX(hsx)', 'MaHang phải tồn tại. Validate TenHang.'),
    ('HangSanXuat', 'XoaHSX(maHang)', 'Không được xóa nếu còn LoaiSanPham đang IsDeleted=0 thuộc hãng đó.'),
    ('LoaiSanPham', 'ThemLoaiSP(lsp)', 'TenLoai không rỗng, ≤200 ký tự. DanhMuc phải là "Laptop"/"Chuột"/"Bàn Phím". ThoiGianBaoHanh > 0. GiaBanGoc ≥ 0. HangSanXuat phải tồn tại.'),
    ('LoaiSanPham', 'XoaLoaiSP(maLoaiSP)', 'Không được xóa nếu còn serial nào có TrangThai ≠ "Đã Bán" thuộc loại này.'),
    ('CauHinh', 'ThemCauHinh(ch)', 'MaCauHinh, MaLoaiSP, TenThuocTinh không rỗng. TenThuocTinh ≤150 ký tự. LoaiSanPham phải tồn tại.'),
    ('SanPham', 'ThemSanPham(sp)', 'MaSerialSP ≤50 ký tự, không rỗng, không trùng (kể cả đã xóa mềm). NgayNhap không trong tương lai. TrangThai mặc định "Trong Kho". Có logic đặc biệt: nếu có serial ảo "x-" đang chờ → tự động thay thế bằng serial thật.'),
    ('SanPham', 'CapNhatTrangThaiSerial(serial, tt)', 'TrangThai phải thuộc {"Trong Kho","Đã Bán","Bảo Hành","Lỗi","Đổi Trả"}. Không cập nhật nếu trạng thái mới = trạng thái cũ.'),
    ('SanPham', 'Xoa(maSerial)', 'Chỉ được xóa mềm serial có TrangThai = "Trong Kho".'),
]
for i, (nhom, method, check) in enumerate(bus_sp_data):
    shade = 'E2EFDA' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [nhom, method, check], shade=shade)
doc.add_paragraph()

# ── BUS_TaiKhoan ──────────────────────────────────
add_heading(doc, '4.2 BUS_TaiKhoan — Xác thực và quản lý tài khoản', 2)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Nhóm', 'Phương Thức', 'Kiểm Tra / Điều Kiện'])
bus_tk_data = [
    ('Mật khẩu', 'HashMatKhau(matKhauGoc)', 'Dùng SHA-256 để hash. Không bao giờ lưu plaintext. Trả về chuỗi hex 64 ký tự.'),
    ('Mật khẩu', 'XacNhanMatKhau(matKhauGoc, hashDaLuu)', 'So sánh hash. Hỗ trợ tương thích ngược với mật khẩu plaintext (kiểm tra trực tiếp trước).'),
    ('NV Login', 'DangNhapNV(tenDN, matKhau, IP)', 'Kiểm tra tài khoản tồn tại → kiểm tra TrangThai ≠ "Khóa" → kiểm tra mật khẩu. Ghi LichSuDangNhap dù thành công hay thất bại.'),
    ('KH Login', 'DangNhapKH(tenDN, matKhau)', 'Tương tự NV nhưng không ghi lịch sử đăng nhập.'),
    ('TK NV', 'ThemTaiKhoanNV(tk)', 'MaTK, MaNV, MaVaiTro, TenDangNhap, MatKhau không rỗng. Mật khẩu ≥6 ký tự. Không trùng TenDangNhap. Tự hash mật khẩu trước khi gọi DAL.'),
    ('TK NV', 'DoiMatKhauNV(maTK, matKhauCu, matKhauMoi)', 'Kiểm tra matKhauCu đúng. matKhauMoi ≥6 ký tự. Hash mật khẩu mới trước khi lưu.'),
    ('TK NV', 'CapNhatTrangThaiNV(maTK, trangThai)', 'TrangThai chỉ nhận "Hoạt Động" hoặc "Khóa".'),
    ('TK KH', 'DangKyTaiKhoanKH(tk)', 'MaKH phải được tạo trước. TenDangNhap không trùng. Khách hàng chưa có tài khoản. Mật khẩu ≥6 ký tự.'),
    ('Sinh mã', 'TaoMaTKNVMoi()', 'Tự động tăng tiến: TKNV000001, TKNV000002, ... Đọc mã lớn nhất từ DB.'),
    ('Sinh mã', 'TaoMaKHLeMoi()', 'KH00000001, KH00000002, ... (prefix KH)'),
    ('Sinh mã', 'TaoMaKHSiMoi()', 'DN00000001, DN00000002, ... (prefix DN — doanh nghiệp)'),
]
for i, (nhom, method, check) in enumerate(bus_tk_data):
    shade = 'FFF2CC' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [nhom, method, check], shade=shade)
doc.add_paragraph()

# ── BUS_KhuyenMai ──────────────────────────────────
add_heading(doc, '4.3 BUS_KhuyenMai — Nghiệp vụ khuyến mãi (data-driven)', 2)
add_para(doc,
    'BUS_KhuyenMai được thiết kế theo nguyên tắc data-driven — logic kiểm tra '
    'điều kiện KM đọc từ CSDL, không hardcode mã KM cụ thể. Mọi chương trình KM '
    'mới thêm vào bảng KhuyenMai đều được xử lý tự động theo cùng một luồng.')
doc.add_paragraph()

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Phần', 'Phương Thức', 'Kiểm Tra / Logic'])
bus_km_data = [
    ('Truy vấn', 'LayDanhSachKhuyenMai()', 'Lấy tất cả KM. Không có điều kiện lọc.'),
    ('Truy vấn', 'LayKhuyenMaiHieuLuc(ngay)', 'Lấy KM đang trong thời gian NgayBatDau ≤ ngay ≤ NgayKetThuc.'),
    ('Nghiệp vụ', 'KiemTraDieuKienKM(km, kh, chiTiet, ngayDat)', '3 bước: (1) Thời gian hiệu lực. (2) Đối tượng: "Tất Cả" / "HSSV" (kiểm tra LaHSSV) / "Doanh Nghiệp" (KH sỉ). (3) Số lượng tối thiểu: đếm serial theo DieuKien (Laptop/Chuột/Bàn Phím) hoặc tổng serial.'),
    ('Nghiệp vụ', 'TinhTienGiam(km, chiTiet)', 'MucGiamSP → giảm % trên từng SP hợp lệ (lọc theo DieuKien nếu có). MucGiamDH → giảm % trên tổng đơn. Chỉ 1 trong 2 được set cho 1 KM.'),
    ('Nghiệp vụ', 'TinhKhuyenMai(dh, chiTiet)', 'Lấy danh sách KM đủ điều kiện, tính tiền giảm từng KM, chọn KM giảm nhiều nhất. Trả về (tienGiam, maKMApDung).'),
    ('CRUD', 'ThemKhuyenMai(km)', 'TenKM không rỗng. DoiTuong phải là "Tất Cả"/"HSSV"/"Doanh Nghiệp". NgayKetThuc ≥ NgayBatDau. MucGiamSP hoặc MucGiamDH phải có ít nhất 1 (không được cả 2). Mức giảm trong 0-100%. DieuKien phải là "Laptop"/"Chuột"/"Bàn Phím" nếu có.'),
    ('CRUD', 'XoaKhuyenMai(maKM)', 'Xóa vật lý. Chỉ xóa được khi chưa có DonHang nào dùng KM này (FK bảo vệ từ DB).'),
]
for i, (phan, method, check) in enumerate(bus_km_data):
    shade = 'FCE4D6' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [phan, method, check], shade=shade)
doc.add_paragraph()

# ── BUS_DonHang ──────────────────────────────────
add_heading(doc, '4.4 BUS_DonHang — Quản lý đơn hàng (giao dịch atomic)', 2)
add_para(doc,
    'Đây là lớp phức tạp nhất — việc tạo/hủy đơn đều là thao tác atomic '
    'chạy trong SQL Transaction để đảm bảo tính nhất quán dữ liệu.')

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Nhóm', 'Phương Thức', 'Kiểm Tra / Logic'])
bus_dh_data = [
    ('Truy vấn', 'LayDanhSachDonHang()', 'Lấy tất cả đơn hàng.'),
    ('Truy vấn', 'LayChiTietDonHang(maDH)', 'Lấy chi tiết đơn hàng kèm thông tin sản phẩm (join LoaiSanPham, HangSanXuat).'),
    ('Tạo đơn', 'TaoDonHang(dh, danhSachSerial, maTK)', 'Validate đầu vào. Kiểm tra KH tồn tại. KH Lẻ: MaHD phải NULL. KH Sỉ: MaHD phải thuộc HopDong "Hiệu Lực". Kiểm tra từng serial: tồn tại, chưa xóa, TrangThai="Trong Kho". Lấy giá từ LoaiSanPham.GiaBanGoc (snapshot). Tính KM. TRANSACTION 3 bước: INSERT DonHang → INSERT ChiTietDonHang → UPDATE SanPham→"Đã Bán".'),
    ('Chuyển TT', 'ChuyenSangDangGiao(maDH)', 'Chỉ từ "Chờ Xử Lý". Đơn thuộc HopDong không được chuyển.'),
    ('Chuyển TT', 'HoanThanhDonHang(maDH)', 'Chỉ từ "Đang Giao". Đơn thuộc HopDong không được chuyển.'),
    ('Hủy đơn', 'HuyDonHang(maDH)', 'Chỉ khi TrangThai = "Chờ Xử Lý". TRANSACTION 2 bước: UPDATE DonHang→"Huỷ" + RESTORE SanPham→"Trong Kho".'),
    ('Tiện ích', 'TaoMaDHMoi()', 'Đọc mã lớn nhất từ DB, cộng 1: DH00000001, DH00000002,...'),
    ('Tìm kiếm', 'TimKiemNhieuDieuKien(...)', 'Lọc phía ứng dụng theo: maKH, tenNV, phuongThuc, maKM, maHD, maDH, maLoaiSP, trangThai. Không có tham số nào bắt buộc.'),
    ('Xem trước', 'XemTruocGiaDon(dh, chiTietTam, maKMChon)', 'Tính tiền giảm trước khi tạo đơn thật. Trả về (tongTien, tienGiam, tienSauGiam, maKMApDung).'),
]
for i, (nhom, method, check) in enumerate(bus_dh_data):
    shade = 'DEEAF1' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [nhom, method, check], shade=shade)
doc.add_paragraph()

# ── BUS_HauMai ──────────────────────────────────
add_heading(doc, '4.5 BUS_HauMai — Bảo hành, đổi trả, khiếu nại', 2)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Module', 'Phương Thức', 'Kiểm Tra / Điều Kiện'])
bus_hm_data = [
    ('Đổi trả', 'KiemTraDieuKienDoiTra(maSerial, ngayMua)',
     '(1) Serial tồn tại. (2) Trong 30 ngày kể từ NgayDat DonHang. (3) Serial chưa có phiếu đổi trả nào (UNIQUE). Trả về tuple (hopLe, lyDoTuChoi).'),
    ('Đổi trả', 'TaoPhieuDoiTra(pdt)',
     'Kiểm tra tất cả trường bắt buộc. LoaiXuLy phải là "Đổi Máy"/"Hoàn Tiền"/"Từ Chối". DonHang phải tồn tại. Gọi KiemTraDieuKienDoiTra. Sau tạo thành công → cập nhật TrangThai SanPham → "Đổi Trả".'),
    ('Đổi trả', 'CapNhatTrangThaiDoiTra(maPDT, tt)', 'TrangThai chỉ nhận: "Đang Xử Lý" / "Hoàn Thành" / "Từ Chối". Phiếu phải tồn tại.'),
    ('Bảo hành', 'TinhNgayKetThucBaoHanh(ngayDat, soThang)', 'ThoiGianBaoHanh > 0. Trả về ngayDat.AddMonths(soThang).'),
    ('Bảo hành', 'TaoPhieuBaoHanh(pbh, loaiKH, soThang)',
     'KH Lẻ: LoaiBH = "Cửa Hàng", NgayKetThuc = NgayDat + ThoiGianBaoHanh. KH Sỉ: LoaiBH = "Hãng", NgayKetThuc do người dùng cung cấp. Kiểm tra NgayKetThuc > NgayBatDau. Kiểm tra serial không có phiếu đang "Đang Xử Lý". Sau tạo thành công → cập nhật SanPham → "Bảo Hành".'),
    ('Bảo hành', 'XuLyTruongHopDacBietKHSi(maDH, maLoaiSP, listSerial)',
     'Điều kiện đặc biệt: trong 30 ngày từ DonHang, ≥10 serial của cùng LoaiSanPham bị lỗi NSX → cửa hàng đổi 1:1, cập nhật TrangThai → "Lỗi".'),
    ('Bảo hành', 'KiemTraConBaoHanh(maSerial)', 'Tìm phiếu BH mới nhất (Đang Xử Lý / Hoàn Thành). Trả về (conBaoHanh, ngayHetHan).'),
    ('Khiếu nại', 'TaoDonKhieuNai(dkn)',
     'MaDH phải liên kết DonHang TrangThai = "Hoàn Thành". MaKH trên đơn khiếu nại phải khớp với DonHang. NoiDung không rỗng. TrangThai khởi tạo = "Đang Xử Lý".'),
]
for i, (mod, method, check) in enumerate(bus_hm_data):
    shade = 'FBE4D5' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [mod, method, check], shade=shade)
doc.add_paragraph()

# ── BUS_BaoCao ──────────────────────────────────
add_heading(doc, '4.6 BUS_BaoCao — Báo cáo thống kê', 2)

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Phương Thức', 'Nội Dung Báo Cáo'])
bao_cao_data = [
    ('BaoCaoTonKho(danhMucLoc?)',
     'Thống kê từng LoaiSanPham: SoLuongTonKho, SoLuongDaBan, SoLuongBaoHanh, TongSoLuong. Lọc theo DanhMuc nếu cần.'),
    ('ThongKeTonKhoTheoDanhMuc()',
     'Tóm tắt tổng hợp theo 3 danh mục (Laptop/Chuột/Bàn Phím): SoLoaiSP, TongTonKho, TongDaBan.'),
    ('BaoCaoNhapHang(tuNgay?, denNgay?, trangThai?)',
     'Danh sách phiếu nhập lọc theo thời gian và/hoặc trạng thái. Validate khoảng ngày, TrangThai phải là "Chờ Xác Nhận"/"Đã Nhập"/"Huỷ".'),
    ('ThongKeNhapHangTheoThang(tuNgay?, denNgay?)',
     'Nhóm theo năm-tháng: SoPhieuNhap, TongTienNhap, TongSoLuongNhap. Chỉ tính phiếu "Đã Nhập".'),
    ('BaoCaoDoanhThu(tuNgay?, denNgay?)',
     'Danh sách đơn hàng TrangThai="Hoàn Thành" lọc theo khoảng thời gian. Validate ngày.'),
    ('ThongKeDoanhThuTheoThang(tuNgay?, denNgay?)',
     'Nhóm theo năm-tháng: SoDonHang, TongDoanhThu, TongDoanhThuSauGiam.'),
    ('ThongKeDoanhThuTheoNam(tuNgay?, denNgay?)',
     'Nhóm theo năm: SoDonHang, TongDoanhThu, TongDoanhThuSauGiam.'),
    ('TinhTongDoanhThu(tuNgay?, denNgay?)',
     'Trả về tuple (tongDoanhThu, tongDoanhThuSauGiam, soDonHang) trong khoảng thời gian.'),
    ('ThongKeTheoHinhThucThanhToan(tuNgay?, denNgay?)',
     'Nhóm theo PhuongThucThanhToan (Tiền Mặt/Chuyển Khoản/Thẻ): SoDonHang, TongDoanhThu.'),
]
for i, (method, desc) in enumerate(bao_cao_data):
    shade = 'E2EFDA' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [method, desc], shade=shade)
doc.add_paragraph()

# ── BUS_KhachHang, BUS_NhanVien, BUS_HopDong, BUS_KhoHang ──
add_heading(doc, '4.7 Các BUS khác', 2)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Class BUS', 'Chức Năng Chính', 'Điều Kiện Nổi Bật'])
other_bus = [
    ('BUS_KhachHang', 'Quản lý KH Lẻ và KH Sỉ: thêm, sửa, xóa mềm, xem chi tiết.',
     'Khi thêm KH Lẻ: transaction: INSERT KhachHang + INSERT KhachHangLe. Không trùng Email. Xóa mềm IsDeleted=1.'),
    ('BUS_NhanVien', 'Quản lý hồ sơ nhân viên: thêm, sửa, xóa mềm.',
     'TenNV, CCCD, SDT không rỗng. Xóa mềm — không xóa vật lý vì liên kết lịch sử đơn hàng.'),
    ('BUS_HopDong', 'Quản lý hợp đồng với KH sỉ: tạo, cập nhật trạng thái, kiểm tra hiệu lực.',
     'KiemTraHopDongCoTheTaoDon(maHD): HopDong phải TrangThai="Hiệu Lực". NgayKy ≤ NgayHetHan.'),
    ('BUS_KhoHang', 'Quản lý nhập kho: PhieuNhap + ChiTietPhieuNhap + NhaCungCap.',
     'Phiếu nhập transaction: INSERT PhieuNhap + INSERT từng ChiTietPhieuNhap + INSERT từng SanPham (serial).'),
]
for i, (bus, chuc, dk) in enumerate(other_bus):
    shade = 'DEEAF1' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [bus, chuc, dk], shade=shade)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHẦN V — GUI WinForm
# ══════════════════════════════════════════════════════════════
add_heading(doc, 'PHẦN V — LỚP GUI (Windows Forms — Ứng dụng nhân viên)', 1)

add_para(doc,
    'GUI_HTQLCuaHangLaptop là ứng dụng Windows Forms dành cho nhân viên '
    'nội bộ. Giao diện MDI (Multiple Document Interface) cho phép mở '
    'nhiều form con đồng thời trong một cửa sổ cha (FormMain). '
    'GUI chỉ gọi BUS — tuyệt đối không gọi DAL trực tiếp.')
doc.add_paragraph()

add_heading(doc, '5.1 Phân quyền truy cập (FormMain)', 2)
add_para(doc,
    'FormMain kiểm tra vai trò người dùng trước khi mở mỗi form con. '
    'Mỗi button chức năng được map với danh sách mã vai trò được phép:')

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Button / Module', 'Vai Trò Được Phép', 'Form Mở Ra'])
pq_data = [
    ('Quản lý hệ thống', 'VT001 (Quản trị)', 'FormQuanLyHeThong'),
    ('Quản lý đơn hàng', 'VT001, VT002 (NVBH), VT004 (CSKH)', 'FormQuanLyDonHang'),
    ('Quản lý hợp đồng', 'VT001, VT002', 'FormQuanLyHopDong'),
    ('Quản lý kho hàng', 'VT001, VT003 (NV Kho), VT004', 'FormQuanLyKhoHang'),
    ('Khuyến mãi', 'VT001, VT002', 'FormQuanLyKhuyenMai'),
    ('Bảo hành', 'VT001, VT004', 'FormBaoHanh'),
    ('Đổi trả', 'VT001, VT004', 'FormDoiTraSanPham'),
    ('Khiếu nại', 'VT001, VT004', 'FormKhieuNai'),
    ('Danh mục sản phẩm', 'VT001, VT003, VT004', 'FormDanhMucSanPham'),
    ('Báo cáo thống kê', 'VT001, VT005 (QL/GĐ)', 'FormBaoCaoThongKe'),
    ('Quản lý khách hàng', 'VT001, VT002, VT004', 'FormQuanLyKhachHang'),
]
for i, (btn, vt, form) in enumerate(pq_data):
    shade = 'E2EFDA' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [btn, vt, form], shade=shade)
doc.add_paragraph()

add_heading(doc, '5.2 Danh sách tất cả Form và chức năng', 2)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Tên Form', 'Chức Năng Chính', 'Điều Kiện / Logic Thêm trong GUI'])
form_data = [
    ('FormDangNhap', 'Nhập tên đăng nhập + mật khẩu để xác thực. Có nút Đổi Mật Khẩu và Thoát.',
     'Kiểm tra không rỗng trước khi gọi BUS. Ẩn mật khẩu bằng PasswordChar. Mở FormThayDoiMatKhau khi click Đổi Mật Khẩu.'),
    ('FormMain', 'MDI Container chính. Hiển thị thông tin nhân viên, vai trò, thời gian đăng nhập. Menu điều hướng.',
     'Kiểm tra phân quyền theo Dictionary mã vai trò → danh sách button. Chuẩn hóa mã vai trò DB 10 ký tự → VT001 format.'),
    ('FormThayDoiMatKhau', 'Đổi mật khẩu tài khoản nhân viên.',
     'Kiểm tra mật khẩu mới ≥6 ký tự, mật khẩu xác nhận phải khớp trước khi gọi BUS.'),
    ('FormQuanLyHeThong', 'Menu con dẫn đến: QuanLyNhanVien, TaiKhoanNhanVien, PhanQuyen.',
     'Chỉ VT001 mới truy cập được.'),
    ('QuanLyNhanVien', 'CRUD nhân viên: Thêm, Sửa, Xóa mềm, Tìm kiếm theo tên/mã.',
     'Kiểm tra ngày sinh hợp lệ (không trong tương lai). Xác nhận trước khi xóa mềm.'),
    ('TaiKhoanNhanVien', 'CRUD tài khoản NV: Thêm TK mới, khóa/mở khóa, đổi mật khẩu.',
     'Chỉ cho phép tạo 1 TK/NV. Tự động sinh mã TKNV. Bind ComboBox vai trò từ BUS.'),
    ('PhanQuyen', 'Xem và cập nhật mô tả vai trò (VT001-VT005).',
     'Chỉ VT001. Không tạo thêm vai trò mới.'),
    ('FormQuanLyKhachHang', 'CRUD khách hàng Lẻ/Sỉ. Tìm kiếm đa tiêu chí. Xem lịch sử mua hàng.',
     'Khi thêm KH Lẻ: bổ sung form thông tin LaHSSV và ngày sinh. Xác nhận xóa mềm.'),
    ('FormQuanLyDonHang', 'Tạo đơn hàng mới. Xem danh sách. Chuyển trạng thái. Hủy đơn.',
     'Tạo đơn: chọn KH → chọn loại SP → hệ thống tự chọn serial → tính KM → xem trước giá → xác nhận. Kiểm tra KH sỉ phải có HopDong. Hiển thị tất cả KM có thể áp dụng.'),
    ('FormQuanLyHopDong', 'Quản lý hợp đồng KH sỉ: tạo, xem, cập nhật trạng thái.',
     'Chỉ tạo HĐ cho KH sỉ. NgayHetHan > NgayKy. Kiểm tra KH tồn tại.'),
    ('FormQuanLyKhoHang', 'Menu con: FormQuanLyNhapHang, FormTonKho, FormQuanLyNhaCungCap, FormDoiTraNCC.', ''),
    ('FormQuanLyNhapHang', 'Tạo phiếu nhập: chọn NCC, chọn LoaiSP, nhập số lượng, nhập serial từng máy.',
     'Kiểm tra NCC tồn tại. Kiểm tra LoaiSP tồn tại. Không trùng serial. Transaction: PhieuNhap + ChiTietPhieuNhap + SanPham.'),
    ('FormTonKho', 'Xem tồn kho theo LoaiSanPham. Lọc theo DanhMuc, TrangThai. Xuất báo cáo.',
     'Hiển thị SoLuongTonKho, DaBan, BaoHanh, Lỗi theo từng LoaiSP.'),
    ('FormQuanLyNhaCungCap', 'CRUD nhà cung cấp. Tìm kiếm, xóa mềm.',
     'Không xóa NCC nếu còn PhieuNhap liên kết chưa hoàn thành.'),
    ('FormDoiTraNCC', 'Xử lý đổi trả serial lỗi cho NCC (trường hợp KH sỉ ≥10 máy lỗi NSX).',
     'Gọi BUS_HauMai.XuLyTruongHopDacBietKHSi. Hiển thị kết quả số serial đã cập nhật trạng thái.'),
    ('FormDanhMucSanPham', 'Menu con: FormHangSanXuat, FormLoaiSanPham, FormCauHinh, FormBaoCaoThongKe tab serial.',
     ''),
    ('FormHangSanXuat', 'CRUD hãng sản xuất. Tìm kiếm, xóa mềm.',
     'Không trùng TenHang. Không xóa nếu còn LoaiSP đang hoạt động.'),
    ('FormLoaiSanPham', 'CRUD loại sản phẩm. Lọc theo hãng, danh mục.',
     'DanhMuc phải là Laptop/Chuột/Bàn Phím. ThoiGianBaoHanh > 0. GiaBanGoc ≥ 0. Không xóa nếu còn serial chưa bán.'),
    ('FormCauHinh', 'Thêm/sửa/xóa cấu hình kỹ thuật cho LoaiSanPham.',
     'TenThuocTinh ≤150 ký tự. LoaiSP phải tồn tại.'),
    ('FormBaoHanh', 'Tạo phiếu bảo hành. Tìm kiếm theo serial/KH. Cập nhật kết quả.',
     'Nhập serial → hệ thống tự tra DonHang, xác định LoaiKH, tính NgayKetThuc. Không tạo nếu có phiếu đang "Đang Xử Lý".'),
    ('FormDoiTraSanPham', 'Tạo phiếu đổi trả. Kiểm tra 30 ngày, chưa có phiếu trước.',
     'Nhập serial → kiểm tra điều kiện → chọn LoaiXuLy → lưu. Sau lưu: TrangThai SanPham → "Đổi Trả".'),
    ('FormKhieuNai', 'Tạo và quản lý đơn khiếu nại. Liên kết DonHang đã "Hoàn Thành".',
     'MaDH phải "Hoàn Thành". MaKH phải khớp DonHang.'),
    ('FormQuanLyKhuyenMai', 'CRUD chương trình khuyến mãi. Xem KM đang hiệu lực.',
     'Validate đầy đủ (xem BUS_KhuyenMai.KiemTraHopLeKM). Hiển thị trạng thái "Đang diễn ra"/"Sắp diễn ra"/"Đã kết thúc".'),
    ('FormBaoCaoThongKe', 'Báo cáo tồn kho, nhập hàng, doanh thu theo khoảng ngày. Xuất dữ liệu.',
     'Lọc khoảng ngày phía GUI. Hiển thị thống kê theo tháng/năm. Tab riêng cho từng loại báo cáo.'),
]
for i, (form, chuc, them) in enumerate(form_data):
    shade = 'DEEAF1' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [form, chuc, them], shade=shade)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHẦN VI — WEB MVC
# ══════════════════════════════════════════════════════════════
add_heading(doc, 'PHẦN VI — LỚP WEB (ASP.NET MVC — Website khách hàng)', 1)

add_para(doc,
    'Website được xây dựng bằng ASP.NET Core MVC, phục vụ khách hàng '
    'mua sắm trực tuyến. Thông tin đăng nhập được lưu trong Session '
    '(timeout 30 phút). Web cũng chỉ gọi BUS — không gọi DAL trực tiếp.')
doc.add_paragraph()

add_heading(doc, '6.1 Danh sách Controllers và chức năng', 2)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Controller', 'Các Action / Route', 'Điều Kiện / Logic Thêm trong Web'])
web_data = [
    ('HomeController\n(HomeController.cs)',
     'Index: Trang chủ, hiển thị sản phẩm nổi bật theo danh mục.',
     'Không cần đăng nhập. Lấy dữ liệu từ BUS_SanPham.'),
    ('SanPhamController\n(SanPhamController.cs)',
     'Index: Danh sách SP lọc theo DanhMuc, hãng, giá.\nChiTiet(maLoaiSP): Chi tiết SP + cấu hình kỹ thuật.',
     'Kiểm tra LoaiSP tồn tại. Hiển thị TonKho (đếm serial "Trong Kho"). Hiển thị CauHinh từ BUS.'),
    ('TaiKhoanController\n(TaiKhoanController.cs)',
     'DangKy GET/POST: Đăng ký tài khoản mới.\nDangNhap GET/POST: Đăng nhập.\nDangXuat POST: Đăng xuất.\nThongTin GET: Xem thông tin cá nhân.\nCapNhatThongTin POST: Cập nhật Email, SDT, DiaChi.\nDoiMatKhau POST: Đổi mật khẩu.',
     'DangKy: tạo KH + KhachHangLe + TaiKhoanKH trong 3 bước. Validate ModelState. DangNhap: Lưu MaKH, MaTK, TenKH, LoaiKH vào Session. DoiMatKhau: matKhauMoi phải khớp xacNhanMatKhauMoi trước khi gọi BUS.'),
    ('GioHangController\n(GioHangController.cs)',
     'Index GET: Xem giỏ hàng.\nThem POST: Thêm LoaiSP vào giỏ.\nXoa POST: Xóa item khỏi giỏ.\nCapNhatSoLuong POST: Cập nhật số lượng.\nDatHang GET: Trang xác nhận đặt hàng.\nXacNhanDatHang POST: Tạo đơn hàng thật.\nThanhToanQR GET: Hiển thị QR code.\nXacNhanQuetQR GET: Xác nhận thanh toán chuyển khoản.\nKiemTraTrangThaiDH GET (JSON API): Kiểm tra trạng thái đơn.',
     'Giỏ hàng lưu trong Session dạng JSON. Them: kiểm tra tồn kho (đếm serial "Trong Kho"), giới hạn SoLuong ≤ TonKho. DatHang: bắt đăng nhập. Lấy danh sách KM đủ điều kiện hiển thị. XacNhanDatHang: tự động chọn serial thật từ kho → gọi BUS_DonHang.TaoDonHang. Nếu chuyển khoản → redirect sang ThanhToanQR + sinh QR code.'),
    ('DonHangController\n(DonHangController.cs)',
     'Index GET: Danh sách đơn hàng của KH đã đăng nhập.\nChiTiet(maDH): Chi tiết đơn.',
     'Chỉ hiển thị đơn của chính KH (so sánh MaKH trong Session). Redirect đăng nhập nếu chưa login.'),
    ('BaoHanhController\n(BaoHanhController.cs)',
     'Index GET/POST: Tra cứu bảo hành theo serial.\nYeuCau GET: Form yêu cầu bảo hành.\nGuiYeuCau POST: Tạo phiếu bảo hành.',
     'Nhập serial → BUS tra thông tin: spTonTai, daBan, conBaoHanh, ngayHetHan. Chỉ tạo phiếu nếu còn trong hạn và đã bán.'),
    ('DoiTraController\n(DoiTraController.cs)',
     'Index GET/POST: Tra cứu điều kiện đổi trả.\nYeuCau GET: Form yêu cầu.\nGuiYeuCau POST: Tạo phiếu đổi trả.',
     'Gọi BUS_HauMai.KiemTraDieuKienDoiTra. Hiển thị hopLe và lyDoTuChoi rõ ràng cho KH.'),
    ('KhieuNaiController\n(KhieuNaiController.cs)',
     'Index GET: Danh sách khiếu nại của KH.\nTao GET: Form tạo khiếu nại.\nGuiKhieuNai POST: Tạo đơn khiếu nại.',
     'MaDH phải là đơn của chính KH và TrangThai="Hoàn Thành". NoiDung không rỗng.'),
]
for i, (ctrl, actions, logic) in enumerate(web_data):
    shade = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [ctrl, actions, logic], shade=shade)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHẦN VII — SO SÁNH GUI và WEB vs BUS
# ══════════════════════════════════════════════════════════════
add_heading(doc, 'PHẦN VII — SO SÁNH GUI/WEB VS BUS: CÁC ĐIỀU KIỆN BỔ SUNG', 1)

add_para(doc,
    'BUS đảm nhận phần lớn logic nghiệp vụ, tuy nhiên GUI và Web có thêm '
    'một số kiểm tra/chức năng ở tầng giao diện để cải thiện trải nghiệm người dùng. '
    'Dưới đây là tổng hợp những gì GUI/Web làm thêm so với BUS:')
doc.add_paragraph()

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
make_table_header(tbl, ['Module', 'BUS đã xử lý', 'GUI thêm', 'Web thêm'])
cmp_data = [
    ('Đăng nhập',
     'Xác thực MK, kiểm tra TrangThai="Khóa", ghi LichSuDangNhap.',
     'Kiểm tra trống tên/MK trước khi gọi BUS. Enter để đăng nhập (AcceptButton).',
     'Redirect về returnUrl sau đăng nhập. Lưu session (MaKH, TenKH, LoaiKH, MaTK).'),
    ('Đăng ký KH',
     'Tạo KH + KhachHangLe + TaiKhoan. Hash MK. Kiểm tra trùng TenDangNhap.',
     'N/A (chỉ có trên Web)',
     'ModelState validation (Data Annotations). Redirect về DangNhap sau thành công.'),
    ('Phân quyền',
     'Không xử lý — BUS chỉ thực thi.',
     'Kiểm tra MaVaiTro trước khi mở từng Form. Thông báo "Không có quyền".',
     'Kiểm tra Session MaKH trước các action yêu cầu đăng nhập. Redirect DangNhap.'),
    ('Tạo đơn hàng',
     'Validate serial, tính KM, TRANSACTION 3 bước.',
     'Giao diện chọn LoaiSP → hiển thị serial available → xem trước giá → xác nhận.',
     'Tự động chọn serial theo số lượng. Kiểm tra đủ tồn kho trước khi đặt. Hỗ trợ QR thanh toán.'),
    ('Giỏ hàng',
     'N/A — BUS không có khái niệm giỏ hàng.',
     'N/A',
     'Giỏ hàng lưu Session JSON. Giới hạn SoLuong ≤ TonKho. Badge số lượng và tổng tiền cập nhật real-time.'),
    ('Thanh toán QR',
     'N/A',
     'N/A',
     'Sinh URL quét QR bằng api.qrserver.com. Khi quét: chuyển đơn "Chờ Xử Lý" → "Đang Giao". Polling trạng thái qua AJAX endpoint KiemTraTrangThaiDH.'),
    ('Đổi trả / Bảo hành',
     'Kiểm tra 30 ngày, chưa có phiếu, cập nhật TrangThai SP.',
     'Nhập serial → auto tra thông tin đơn → hiển thị kết quả kiểm tra. Xác nhận trước khi tạo phiếu.',
     'Form tra cứu serial → hiển thị hopLe/lyDoTuChoi. Form yêu cầu riêng biệt.'),
    ('Báo cáo',
     'Tính toán, lọc, tổng hợp dữ liệu trả về DataTable.',
     'Hiển thị DataTable trong DataGridView. Lọc theo ngày phía GUI. Nhiều Tab báo cáo.',
     'N/A — Báo cáo chỉ dành cho nhân viên (GUI).'),
    ('Sinh mã tự động',
     'TaoMaDHMoi(), TaoMaTKNVMoi(), TaoMaKHLeMoi(), TaoMaKHSiMoi().',
     'Gọi BUS để lấy mã mới trước khi hiển thị form.',
     'Gọi BUS để lấy mã mới khi tạo KH và TK.'),
]
for i, (mod, bus_do, gui_do, web_do) in enumerate(cmp_data):
    shade = 'FFF2CC' if i % 2 == 0 else 'FFFFFF'
    add_table_row(tbl, [mod, bus_do, gui_do, web_do], shade=shade)
doc.add_paragraph()

# ──── KẾT LUẬN ────
add_heading(doc, 'KẾT LUẬN', 1)
add_para(doc,
    'Hệ thống Quản Lý Cửa Hàng Phân Phối Máy Tính được xây dựng theo mô hình '
    'N-Layer đảm bảo sự tách biệt rõ ràng giữa các tầng:')
conclusions = [
    'DTO: "Hộp dữ liệu" — không có logic, chỉ mang dữ liệu.',
    'DAL: "Cầu nối CSDL" — CRUD thuần, SQL tập trung tại đây, bảo vệ SQL Injection.',
    'BUS: "Bộ não nghiệp vụ" — validate dữ liệu, áp dụng quy tắc, điều phối transaction.',
    'GUI: "Giao diện nhân viên" — phân quyền, UX tốt, không gọi thẳng CSDL.',
    'Web: "Giao diện khách hàng" — mua sắm online, giỏ hàng, thanh toán QR.',
]
for c in conclusions:
    add_bullet(doc, c)

doc.add_paragraph()
add_para(doc,
    'Kiến trúc này giúp hệ thống dễ bảo trì, dễ mở rộng — khi cần thêm '
    'chức năng chỉ cần thêm class BUS/DAL mới mà không ảnh hưởng đến GUI '
    'hay Web đang hoạt động. Mọi thay đổi nghiệp vụ tập trung tại BUS, '
    'mọi thay đổi cơ sở dữ liệu tập trung tại DAL.', italic=True)

# ──── LƯU FILE ────
output_path = r'c:\Application\University\3.2\LaptrinhCSDL\LTCSDL_Tieuluan_HethongphanphoiMT\LTCSDL_HTQLCuaHangPhanPhoiLaptop\GiaiThichCode_HTQLCuaHangLaptop.docx'
doc.save(output_path)
print(f'✅ Đã tạo file Word thành công: {output_path}')
